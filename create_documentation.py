#!/usr/bin/env python3
"""
Script to create combined documentation, API documentation, and demo script
for the Cafeteria Pre-ordering System
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_combined_documentation():
    """Create combined Problem Statement + Architecture document"""
    doc = Document()

    # Title
    title = doc.add_heading('Cafeteria Pre-ordering System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Documentation - Problem Statement, Design & Architecture')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # TABLE OF CONTENTS
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        'Part 1: Problem Statement & Design',
        '   1.1 Problem Statement',
        '   1.2 System Design',
        '   1.3 Functional Requirements',
        '   1.4 Non-Functional Requirements',
        '   1.5 User Roles',
        'Part 2: Architecture Documentation',
        '   2.1 Microservices Architecture',
        '   2.2 Service Responsibilities',
        '   2.3 Database Schema',
        '   2.4 Component Hierarchy',
        '   2.5 Communication Patterns',
        'Part 3: Implementation Details',
        '   3.1 Tech Stack',
        '   3.2 Security Implementation',
        '   3.3 Deployment Architecture'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # PART 1: PROBLEM STATEMENT
    doc.add_heading('Part 1: Problem Statement & Design', 1)

    doc.add_heading('1.1 Problem Statement', 2)
    doc.add_paragraph(
        'The Cafeteria Pre-ordering System addresses the challenge of long queues and wait times '
        'in university/institution cafeterias. Students and staff can pre-order their meals, '
        'select pickup times, and pay digitally, while canteen staff can manage orders efficiently.'
    )

    doc.add_heading('1.2 Problem Context', 2)
    doc.add_paragraph(
        'Traditional cafeteria operations face several challenges:'
    )
    challenges = [
        'Long queues during peak hours (lunch breaks)',
        'Uncertainty about food availability',
        'Cash handling and payment delays',
        'No visibility into dietary restrictions/allergens',
        'Manual order tracking and management',
        'Food wastage due to over-preparation'
    ]
    for challenge in challenges:
        doc.add_paragraph(challenge, style='List Bullet')

    doc.add_heading('1.3 Proposed Solution', 2)
    doc.add_paragraph(
        'A full-stack web application with microservices architecture enabling:'
    )
    solutions = [
        'Pre-ordering with scheduled pickup times',
        'Digital wallet for cashless payments',
        'Real-time menu browsing with dietary filters',
        'Order status tracking (pending → confirmed → preparing → ready → completed)',
        'Admin dashboard for menu and order management'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    doc.add_heading('1.4 Functional Requirements', 2)

    doc.add_heading('Student/User Features:', 3)
    user_features = [
        'User registration and authentication with JWT',
        'Browse menu items with category filters',
        'Search menu items by name',
        'View dietary tags (vegetarian, vegan, gluten-free) and allergens',
        'Add items to cart with quantity selection',
        'Checkout with pickup date and time selection',
        'Digital wallet for payments (add funds, view balance)',
        'Track order status in real-time',
        'View order history',
        'Update profile with dietary preferences and allergies',
        'Cancel pending orders'
    ]
    for feature in user_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Admin/Canteen Features:', 3)
    admin_features = [
        'Admin dashboard with statistics (total orders, pending orders, revenue)',
        'CRUD operations for menu items',
        'Update order status through workflow',
        'View all orders with filtering',
        'Manage item availability and daily order limits',
        'Seed sample data for testing'
    ]
    for feature in admin_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('1.5 Non-Functional Requirements', 2)
    nfrs = [
        'Microservices architecture for scalability',
        'Database per service pattern',
        'JWT-based authentication',
        'Role-based access control',
        'Responsive UI for mobile and desktop',
        'API documentation via Swagger/OpenAPI',
        'PostgreSQL for data persistence'
    ]
    for nfr in nfrs:
        doc.add_paragraph(nfr, style='List Bullet')

    doc.add_heading('1.6 User Roles', 2)
    roles_table = doc.add_table(rows=4, cols=3)
    roles_table.style = 'Light Grid Accent 1'
    hdr_cells = roles_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Permissions'
    hdr_cells[2].text = 'Typical Users'

    row1 = roles_table.rows[1].cells
    row1[0].text = 'Student'
    row1[1].text = 'Browse menu, place orders, manage profile, track orders'
    row1[2].text = 'Students'

    row2 = roles_table.rows[2].cells
    row2[0].text = 'Canteen'
    row2[1].text = 'Update order status, view orders'
    row2[2].text = 'Canteen staff'

    row3 = roles_table.rows[3].cells
    row3[0].text = 'Admin'
    row3[1].text = 'Full menu management, all order operations, dashboard'
    row3[2].text = 'Administrators'

    doc.add_page_break()

    # PART 2: ARCHITECTURE
    doc.add_heading('Part 2: Architecture Documentation', 1)

    doc.add_heading('2.1 Microservices Architecture', 2)
    doc.add_paragraph(
        'The system follows a microservices architecture pattern with the following components:'
    )

    doc.add_heading('Architecture Diagram (Text Representation):', 3)
    arch_text = """
    ┌─────────────────────────────────────────────────────────────────────────────┐
    │                              Client Layer                                    │
    │                        React + TypeScript                                    │
    ├─────────────────────────────────────────────────────────────────────────────┤
    │                        API Gateway (Port 5000)                               │
    │  ┌─────────────────────────────────────────────────────────────────────┐    │
    │  │   Ocelot Gateway - Routing, Load Balancing, Authentication           │    │
    │  │   ┌──────────────┐  ┌──────────────┐  ┌──────────────┐              │    │
    │  │   │  Auth Routes │  │  Menu Routes │  │ Order Routes │              │    │
    │  │   └──────┬───────┘  └──────┬───────┘  └──────┬───────┘              │    │
    │  └──────────┼────────────────┼────────────────┼───────────────────────┘    │
    └─────────────┼────────────────┼────────────────┼──────────────────────────────┘
                  │                │                │
                  ▼                ▼                ▼
    ┌─────────────────────────────────────────────────────────────────────────────┐
    │                      Microservices Layer                                     │
    ├─────────────────┬─────────────────┬───────────────────────────────────────┤
    │  Auth Service   │  Menu Service   │  Order Service                        │
    │  (Port 5001)    │  (Port 5002)    │  (Port 5003)                          │
    ├─────────────────┼─────────────────┼───────────────────────────────────────┤
    │  • User Auth    │  • Menu CRUD    │  • Order Management                   │
    │  • JWT Tokens   │  • Categories   │  • Status Workflow                    │
    │  • User Profile │  • Filtering    │  • Order Items                        │
    │  • Wallet       │  • Search       │  • Statistics                         │
    └────────┬────────┴────────┬────────┴────────┬──────────────────────────────┘
             │                   │                   │
             │   ┌───────────────┴───────────────────┘
             │   │
             ▼   ▼
    ┌─────────────────────────────────────────────────────────────┐
    │                    PostgreSQL Databases                      │
    ├─────────────────┬─────────────────┬─────────────────────────┤
    │ cafeteria_auth  │ cafeteria_menu  │ cafeteria_orders        │
    │                 │                 │                         │
    │ • users table   │ • menu_items    │ • orders                │
    │                 │                 │ • order_items           │
    └─────────────────┴─────────────────┴─────────────────────────┘
    """
    para = doc.add_paragraph()
    para.add_run(arch_text).font.name = 'Courier New'

    doc.add_heading('2.2 Service Responsibilities', 2)

    services = [
        ('API Gateway (Port 5000)', [
            'Technology: Ocelot (.NET 10)',
            'Request routing to appropriate microservices',
            'JWT authentication validation',
            'CORS handling',
            'Claims extraction and forwarding (X-User-Id, X-User-Role headers)'
        ]),
        ('Auth Service (Port 5001)', [
            'Database: cafeteria_auth',
            'User registration and login',
            'JWT token generation',
            'User profile management (dietary preferences, allergies)',
            'Wallet management (balance, add funds)'
        ]),
        ('Menu Service (Port 5002)', [
            'Database: cafeteria_menu',
            'Menu item CRUD operations',
            'Category management',
            'Dietary tags and allergens',
            'Availability tracking'
        ]),
        ('Order Service (Port 5003)', [
            'Database: cafeteria_orders',
            'Order creation and management',
            'Order status workflow (pending → confirmed → preparing → ready → completed)',
            'Order statistics',
            'Order history'
        ])
    ]

    for service_name, responsibilities in services:
        doc.add_heading(service_name, 3)
        for resp in responsibilities:
            doc.add_paragraph(resp, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('2.3 Database Schema', 2)

    doc.add_heading('Auth Service (cafeteria_auth)', 3)
    doc.add_paragraph('Table: users')
    users_schema = """
    CREATE TABLE users (
        id SERIAL PRIMARY KEY,
        name VARCHAR(100) NOT NULL,
        email VARCHAR(100) UNIQUE NOT NULL,
        password_hash VARCHAR(255) NOT NULL,
        role VARCHAR(20) DEFAULT 'student',
        dietary_preferences TEXT[],
        allergies TEXT[],
        wallet_balance DECIMAL(18,2) DEFAULT 0.00,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    """
    para = doc.add_paragraph()
    para.add_run(users_schema).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('Menu Service (cafeteria_menu)', 3)
    doc.add_paragraph('Table: menu_items')
    menu_schema = """
    CREATE TABLE menu_items (
        id SERIAL PRIMARY KEY,
        name VARCHAR(100) NOT NULL,
        description TEXT,
        price DECIMAL(18,2) NOT NULL,
        category VARCHAR(50),
        dietary_tags TEXT[],
        allergens TEXT[],
        available BOOLEAN DEFAULT true,
        preparation_time INTEGER DEFAULT 15,
        max_orders_per_day INTEGER DEFAULT 100,
        orders_today INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    """
    para = doc.add_paragraph()
    para.add_run(menu_schema).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('Order Service (cafeteria_orders)', 3)
    doc.add_paragraph('Table: orders')
    orders_schema = """
    CREATE TABLE orders (
        id SERIAL PRIMARY KEY,
        user_id INTEGER NOT NULL,
        total_amount DECIMAL(18,2) NOT NULL,
        pickup_time TIMESTAMP NOT NULL,
        pickup_date DATE NOT NULL,
        status VARCHAR(20) DEFAULT 'pending',
        payment_status VARCHAR(20) DEFAULT 'unpaid',
        special_instructions TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP
    );
    """
    para = doc.add_paragraph()
    para.add_run(orders_schema).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_paragraph('Table: order_items')
    order_items_schema = """
    CREATE TABLE order_items (
        id SERIAL PRIMARY KEY,
        order_id INTEGER REFERENCES orders(id) ON DELETE CASCADE,
        menu_item_id INTEGER NOT NULL,
        quantity INTEGER NOT NULL,
        price DECIMAL(18,2) NOT NULL,
        menu_item_name VARCHAR(100)
    );
    """
    para = doc.add_paragraph()
    para.add_run(order_items_schema).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('2.4 Component Hierarchy', 2)
    doc.add_heading('Frontend Component Tree:', 3)
    component_tree = """
    App
    ├── AuthProvider (Context)
    │   └── Manages global authentication state
    │
    ├── Router
    │   ├── Navbar
    │   │   ├── Logo/Brand
    │   │   ├── Navigation Links
    │   │   ├── User Menu
    │   │   └── Logout Button
    │   │
    │   └── Routes
    │       ├── Public Routes
    │       │   ├── / (Home) → Home
    │       │   ├── /login → Login
    │       │   └── /register → Register
    │       │
    │       └── Protected Routes
    │           ├── /menu → Menu
    │           ├── /cart → Cart
    │           ├── /orders → Orders
    │           ├── /profile → Profile
    │           └── /admin → AdminPanel
    """
    para = doc.add_paragraph()
    para.add_run(component_tree).font.name = 'Courier New'

    doc.add_heading('Key Components:', 3)
    components = [
        ('Layout Components', ['Navbar', 'AuthProvider', 'ProtectedRoute']),
        ('Page Components', ['Home', 'Login', 'Register', 'Menu', 'Cart', 'Orders', 'Profile', 'AdminPanel']),
        ('Shared Components', ['MenuCard', 'CartItem', 'OrderCard', 'StatusBadge'])
    ]
    for category, items in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{category}: ').bold = True
        p.add_run(', '.join(items))

    doc.add_heading('2.5 Communication Patterns', 2)

    doc.add_heading('Synchronous (REST API):', 3)
    doc.add_paragraph('Frontend → API Gateway → Microservices')

    doc.add_heading('Data Isolation:', 3)
    isolation = [
        'Each service owns its database',
        'No direct database access between services',
        'Service-to-service communication via APIs'
    ]
    for item in isolation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('API Gateway Routes:', 3)
    routes_table = doc.add_table(rows=5, cols=4)
    routes_table.style = 'Light Grid Accent 1'
    hdr = routes_table.rows[0].cells
    hdr[0].text = 'Route'
    hdr[1].text = 'Service'
    hdr[2].text = 'Port'
    hdr[3].text = 'Auth Required'

    routes_data = [
        ('/api/auth/*', 'Auth Service', '5001', 'No*'),
        ('/api/users/*', 'Auth Service', '5001', 'Yes'),
        ('/api/menu/*', 'Menu Service', '5002', 'Varies'),
        ('/api/orders/*', 'Order Service', '5003', 'Yes')
    ]

    for i, (route, service, port, auth) in enumerate(routes_data, 1):
        row = routes_table.rows[i].cells
        row[0].text = route
        row[1].text = service
        row[2].text = port
        row[3].text = auth

    doc.add_paragraph('*Except /api/auth/me which requires auth')

    doc.add_page_break()

    # PART 3: IMPLEMENTATION
    doc.add_heading('Part 3: Implementation Details', 1)

    doc.add_heading('3.1 Tech Stack', 2)

    tech_table = doc.add_table(rows=3, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    tech_table.rows[0].cells[0].text = 'Layer'
    tech_table.rows[0].cells[1].text = 'Technologies'

    tech_data = [
        ('Microservices (Backend)', '.NET 10, Ocelot, Entity Framework Core, PostgreSQL, JWT, BCrypt'),
        ('Frontend', 'React 19, TypeScript, React Router, Axios, CSS Modules')
    ]
    for i, (layer, tech) in enumerate(tech_data, 1):
        tech_table.rows[i].cells[0].text = layer
        tech_table.rows[i].cells[1].text = tech

    doc.add_heading('3.2 Security Implementation', 2)
    security = [
        'Authentication: JWT tokens validated at Gateway',
        'Authorization: Role-based checks using X-User-Role header',
        'Data Isolation: Service-level database isolation',
        'CORS: Configured at Gateway level',
        'HTTPS: Production deployments use HTTPS',
        'Password Hashing: BCrypt for secure password storage'
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.3 Deployment Architecture', 2)

    doc.add_heading('Development Setup:', 3)
    dev_arch = """
    ┌─────────────────────────────────────────────────────────┐
    │                    Development Machine                   │
    ├─────────────────────────────────────────────────────────┤
    │  React Frontend (localhost:3000)                        │
    ├─────────────────────────────────────────────────────────┤
    │  API Gateway (localhost:5010)                           │
    ├─────────────────────────────────────────────────────────┤
    │  Auth Service (localhost:5001)                         │
    │  Menu Service (localhost:5002)                         │
    │  Order Service (localhost:5003)                        │
    ├─────────────────────────────────────────────────────────┤
    │  PostgreSQL (localhost:5432)                            │
    │  ├── cafeteria_auth                                     │
    │  ├── cafeteria_menu                                     │
    │  └── cafeteria_orders                                   │
    └─────────────────────────────────────────────────────────┘
    """
    para = doc.add_paragraph()
    para.add_run(dev_arch).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('3.4 Benefits of Microservices Architecture', 2)
    benefits = [
        'Independent Deployment: Each service can be deployed separately',
        'Technology Diversity: Different services can use different tech stacks',
        'Fault Isolation: Failure in one service doesn\'t affect others',
        'Team Autonomy: Different teams can work on different services',
        'Database Per Service: No shared database coupling'
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    # Save document
    output_path = '/Users/subhodipanwesa/Documents/BITS_WILP/FullStack/Assignment1/Complete_Architecture_Documentation.docx'
    doc.save(output_path)
    print(f"✅ Combined documentation created: {output_path}")
    return output_path


def create_api_documentation():
    """Create comprehensive API documentation"""
    doc = Document()

    # Title
    title = doc.add_heading('Cafeteria Pre-ordering System - API Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('REST API Reference Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph('Base URL: http://localhost:5010 (via API Gateway)')
    doc.add_paragraph('Authentication: Bearer Token (JWT)')

    doc.add_page_break()

    # TABLE OF CONTENTS
    doc.add_heading('Table of Contents', 1)
    toc = [
        'Authentication Endpoints',
        'User Endpoints',
        'Menu Endpoints',
        'Order Endpoints',
        'Error Responses',
        'Data Models'
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # AUTHENTICATION ENDPOINTS
    doc.add_heading('1. Authentication Endpoints', 1)

    # Register
    doc.add_heading('1.1 Register User', 2)
    doc.add_paragraph('POST /api/auth/register')
    doc.add_paragraph('Creates a new user account.')

    doc.add_heading('Request Body:', 3)
    request_example = """{
  "name": "John Doe",
  "email": "john@student.com",
  "password": "password123",
  "role": "student"
}"""
    para = doc.add_paragraph()
    para.add_run(request_example).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('Response (201 Created):', 3)
    response_example = """{
  "id": 1,
  "name": "John Doe",
  "email": "john@student.com",
  "role": "student",
  "walletBalance": 0.00,
  "dietaryPreferences": [],
  "allergies": [],
  "createdAt": "2026-05-09T10:00:00Z"
}"""
    para = doc.add_paragraph()
    para.add_run(response_example).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    # Login
    doc.add_heading('1.2 Login', 2)
    doc.add_paragraph('POST /api/auth/login')
    doc.add_paragraph('Authenticates user and returns JWT token.')

    doc.add_heading('Request Body:', 3)
    login_req = """{
  "email": "john@student.com",
  "password": "password123"
}"""
    para = doc.add_paragraph()
    para.add_run(login_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('Response (200 OK):', 3)
    login_resp = """{
  "token": "eyJhbGciOiJIUzI1NiIs...",
  "user": {
    "id": 1,
    "name": "John Doe",
    "email": "john@student.com",
    "role": "student",
    "walletBalance": 100.00
  }
}"""
    para = doc.add_paragraph()
    para.add_run(login_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    # Get Current User
    doc.add_heading('1.3 Get Current User', 2)
    doc.add_paragraph('GET /api/auth/me')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_paragraph('Returns the currently authenticated user.')

    doc.add_page_break()

    # USER ENDPOINTS
    doc.add_heading('2. User Endpoints', 1)

    doc.add_heading('2.1 Get Wallet Balance', 2)
    doc.add_paragraph('GET /api/users/wallet')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Response:', 3)
    wallet_resp = """{
  "balance": 100.00
}"""
    para = doc.add_paragraph()
    para.add_run(wallet_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('2.2 Add Funds to Wallet', 2)
    doc.add_paragraph('POST /api/users/wallet/add')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Request Body:', 3)
    add_funds_req = """{
  "amount": 50.00
}"""
    para = doc.add_paragraph()
    para.add_run(add_funds_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('2.3 Get User Preferences', 2)
    doc.add_paragraph('GET /api/users/preferences')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Response:', 3)
    prefs_resp = """{
  "dietaryPreferences": ["vegetarian", "gluten-free"],
  "allergies": ["peanuts", "shellfish"]
}"""
    para = doc.add_paragraph()
    para.add_run(prefs_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('2.4 Update User Preferences', 2)
    doc.add_paragraph('PUT /api/users/preferences')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Request Body:', 3)
    update_prefs_req = """{
  "dietaryPreferences": ["vegetarian"],
  "allergies": ["peanuts"]
}"""
    para = doc.add_paragraph()
    para.add_run(update_prefs_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # MENU ENDPOINTS
    doc.add_heading('3. Menu Endpoints', 1)

    doc.add_heading('3.1 Get All Menu Items', 2)
    doc.add_paragraph('GET /api/menu')
    doc.add_paragraph('Query Parameters: category, search, dietaryTags, available')
    doc.add_heading('Response:', 3)
    menu_resp = """[
  {
    "id": 1,
    "name": "Veggie Burger",
    "description": "Plant-based patty with fresh vegetables",
    "price": 8.99,
    "category": "Main",
    "dietaryTags": ["vegetarian", "vegan"],
    "allergens": ["gluten", "soy"],
    "available": true,
    "preparationTime": 15,
    "maxOrdersPerDay": 100,
    "ordersToday": 45
  }
]"""
    para = doc.add_paragraph()
    para.add_run(menu_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('3.2 Get Menu Item by ID', 2)
    doc.add_paragraph('GET /api/menu/{id}')
    doc.add_paragraph('Returns single menu item details.')

    doc.add_heading('3.3 Create Menu Item (Admin)', 2)
    doc.add_paragraph('POST /api/menu')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')
    doc.add_heading('Request Body:', 3)
    create_menu_req = """{
  "name": "Grilled Chicken Salad",
  "description": "Fresh salad with grilled chicken",
  "price": 12.99,
  "category": "Salads",
  "dietaryTags": ["high-protein", "gluten-free"],
  "allergens": [],
  "available": true,
  "preparationTime": 10,
  "maxOrdersPerDay": 50
}"""
    para = doc.add_paragraph()
    para.add_run(create_menu_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('3.4 Update Menu Item (Admin)', 2)
    doc.add_paragraph('PUT /api/menu/{id}')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')

    doc.add_heading('3.5 Delete Menu Item (Admin)', 2)
    doc.add_paragraph('DELETE /api/menu/{id}')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')

    doc.add_heading('3.6 Get Categories', 2)
    doc.add_paragraph('GET /api/menu/categories')
    doc.add_heading('Response:', 3)
    cat_resp = """[
  "Main",
  "Beverages",
  "Snacks",
  "Salads",
  "Desserts"
]"""
    para = doc.add_paragraph()
    para.add_run(cat_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('3.7 Seed Sample Data', 2)
    doc.add_paragraph('POST /api/menu/seed')
    doc.add_paragraph('Populates database with sample menu items.')

    doc.add_page_break()

    # ORDER ENDPOINTS
    doc.add_heading('4. Order Endpoints', 1)

    doc.add_heading('4.1 Get All Orders (Admin)', 2)
    doc.add_paragraph('GET /api/orders')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')
    doc.add_paragraph('Query Parameters: status, startDate, endDate')

    doc.add_heading('4.2 Get My Orders', 2)
    doc.add_paragraph('GET /api/orders/my-orders')
    doc.add_paragraph('Authorization: Bearer {token}')

    doc.add_heading('4.3 Get Order by ID', 2)
    doc.add_paragraph('GET /api/orders/{id}')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Response:', 3)
    order_resp = """{
  "id": 1,
  "userId": 1,
  "totalAmount": 25.97,
  "pickupTime": "12:30",
  "pickupDate": "2026-05-09",
  "status": "confirmed",
  "paymentStatus": "paid",
  "specialInstructions": "Extra sauce on the side",
  "items": [
    {
      "id": 1,
      "menuItemId": 1,
      "quantity": 2,
      "price": 8.99,
      "menuItemName": "Veggie Burger"
    },
    {
      "id": 2,
      "menuItemId": 3,
      "quantity": 1,
      "price": 7.99,
      "menuItemName": "Fresh Juice"
    }
  ],
  "createdAt": "2026-05-09T09:00:00Z",
  "updatedAt": "2026-05-09T09:05:00Z"
}"""
    para = doc.add_paragraph()
    para.add_run(order_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('4.4 Create Order', 2)
    doc.add_paragraph('POST /api/orders')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_heading('Request Body:', 3)
    create_order_req = """{
  "items": [
    {
      "menuItemId": 1,
      "quantity": 2,
      "price": 8.99
    },
    {
      "menuItemId": 3,
      "quantity": 1,
      "price": 7.99
    }
  ],
  "pickupTime": "12:30",
  "pickupDate": "2026-05-09",
  "specialInstructions": "Extra sauce on the side"
}"""
    para = doc.add_paragraph()
    para.add_run(create_order_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('4.5 Update Order Status (Admin)', 2)
    doc.add_paragraph('PUT /api/orders/{id}/status')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')
    doc.add_heading('Request Body:', 3)
    status_req = """{
  "status": "preparing"
}"""
    para = doc.add_paragraph()
    para.add_run(status_req).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)
    doc.add_paragraph('Valid statuses: pending, confirmed, preparing, ready, completed, cancelled')

    doc.add_heading('4.6 Cancel Order', 2)
    doc.add_paragraph('DELETE /api/orders/{id}')
    doc.add_paragraph('Authorization: Bearer {token}')
    doc.add_paragraph('Only pending orders can be cancelled.')

    doc.add_heading('4.7 Get Order Statistics', 2)
    doc.add_paragraph('GET /api/orders/stats')
    doc.add_paragraph('Authorization: Bearer {token} (Admin only)')
    doc.add_heading('Response:', 3)
    stats_resp = """{
  "totalOrders": 150,
  "pendingOrders": 12,
  "completedOrders": 130,
  "cancelledOrders": 8,
  "totalRevenue": 2450.50
}"""
    para = doc.add_paragraph()
    para.add_run(stats_resp).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ERROR RESPONSES
    doc.add_heading('5. Error Responses', 1)

    error_table = doc.add_table(rows=8, cols=3)
    error_table.style = 'Light Grid Accent 1'
    hdr = error_table.rows[0].cells
    hdr[0].text = 'Status Code'
    hdr[1].text = 'Meaning'
    hdr[2].text = 'Example'

    errors = [
        ('400', 'Bad Request', 'Invalid request data'),
        ('401', 'Unauthorized', 'Missing or invalid token'),
        ('403', 'Forbidden', 'Insufficient permissions'),
        ('404', 'Not Found', 'Resource not found'),
        ('409', 'Conflict', 'User already exists'),
        ('422', 'Unprocessable', 'Validation error'),
        ('500', 'Server Error', 'Internal server error')
    ]

    for i, (code, meaning, example) in enumerate(errors, 1):
        row = error_table.rows[i].cells
        row[0].text = code
        row[1].text = meaning
        row[2].text = example

    doc.add_heading('Error Response Format:', 2)
    error_format = """{
  "message": "Error description",
  "errors": {
    "fieldName": ["Error message"]
  }
}"""
    para = doc.add_paragraph()
    para.add_run(error_format).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # DATA MODELS
    doc.add_heading('6. Data Models', 1)

    doc.add_heading('6.1 User Model', 2)
    user_model = """{
  "id": integer (read-only),
  "name": string (required, max 100 chars),
  "email": string (required, unique, valid email),
  "password": string (required, min 6 chars, write-only),
  "role": enum ['student', 'staff', 'admin', 'canteen'] (default: 'student'),
  "dietaryPreferences": array of strings,
  "allergies": array of strings,
  "walletBalance": decimal (read-only),
  "createdAt": datetime (read-only)
}"""
    para = doc.add_paragraph()
    para.add_run(user_model).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('6.2 MenuItem Model', 2)
    menu_model = """{
  "id": integer (read-only),
  "name": string (required, max 100 chars),
  "description": string (optional),
  "price": decimal (required, positive),
  "category": string (optional),
  "dietaryTags": array of strings,
  "allergens": array of strings,
  "available": boolean (default: true),
  "preparationTime": integer (minutes, default: 15),
  "maxOrdersPerDay": integer (default: 100),
  "ordersToday": integer (read-only),
  "createdAt": datetime (read-only)
}"""
    para = doc.add_paragraph()
    para.add_run(menu_model).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('6.3 Order Model', 2)
    order_model = """{
  "id": integer (read-only),
  "userId": integer (read-only),
  "totalAmount": decimal (read-only, calculated),
  "pickupTime": string (required, format: HH:mm),
  "pickupDate": date (required, format: YYYY-MM-DD),
  "status": enum ['pending', 'confirmed', 'preparing', 'ready', 'completed', 'cancelled'],
  "paymentStatus": enum ['unpaid', 'paid', 'refunded'],
  "specialInstructions": string (optional),
  "items": array of OrderItem,
  "createdAt": datetime (read-only),
  "updatedAt": datetime (read-only)
}"""
    para = doc.add_paragraph()
    para.add_run(order_model).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    doc.add_heading('6.4 OrderItem Model', 2)
    order_item_model = """{
  "id": integer (read-only),
  "orderId": integer (read-only),
  "menuItemId": integer (required),
  "quantity": integer (required, min 1),
  "price": decimal (required, price at time of order),
  "menuItemName": string (snapshot of item name at order time)
}"""
    para = doc.add_paragraph()
    para.add_run(order_item_model).font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)

    # Save document
    output_path = '/Users/subhodipanwesa/Documents/BITS_WILP/FullStack/Assignment1/API_Documentation.docx'
    doc.save(output_path)
    print(f"✅ API documentation created: {output_path}")
    return output_path


def create_demo_script():
    """Create demo video script markdown file"""
    script = """# Cafeteria Pre-ordering System - Demo Video Script

**Duration:** 8-10 minutes
**Target:** BITS WILP Faculty Evaluation
**Presenter:** Student Name
**Project:** Cafeteria Food Pre-ordering System (Microservices Architecture)

---

## Video Structure

### **INTRO (0:00 - 0:30)**
**[Screen: Title Slide]**

**Narration:**
"Hello, this is the demonstration video for the Cafeteria Food Pre-ordering System, developed as part of the Full Stack Application Development course assignment. This system allows students and staff to pre-order meals, select pickup times, and make digital payments, while enabling canteen staff to manage orders efficiently through a microservices-based architecture."

**[Transition to Problem Statement]**

---

### **SECTION 1: Problem Statement & Solution (0:30 - 1:30)**
**[Screen: Problem Statement Slide or Diagram]**

**Narration:**
"The problem we address is the long queues and wait times during peak hours in institutional cafeterias. Students often face uncertainty about food availability and must wait in line with cash payments. Our solution provides a digital platform for pre-ordering with scheduled pickups, digital wallet payments, and real-time order tracking."

**Key Points to Highlight:**
- Long queues during lunch breaks
- Uncertainty about food availability
- Cash handling delays
- Food wastage due to over-preparation
- Solution benefits: Time-saving, convenience, efficiency

**[Transition to Architecture Overview]**

---

### **SECTION 2: Architecture Overview (1:30 - 3:00)**
**[Screen: Architecture Diagram]**

**Narration:**
"Our system follows a microservices architecture pattern. The frontend is built with React and TypeScript. All requests go through an API Gateway using Ocelot, which handles routing and authentication. Behind the gateway, we have three independent microservices:"

**Walk through each service:**
1. **Auth Service** (Port 5001): "Handles user registration, login, JWT token generation, and wallet management."
2. **Menu Service** (Port 5002): "Manages menu items, categories, dietary tags, and availability tracking."
3. **Order Service** (Port 5003): "Handles order creation, status workflow, and order statistics."

**Emphasize:**
- Database per service pattern (3 PostgreSQL databases)
- JWT validation at Gateway level
- Independent service deployment
- Claims forwarding via headers

**[Transition to Live Demo]**

---

### **SECTION 3: Application Demo - User Features (3:00 - 6:00)**

#### **3.1 Landing Page & Registration (3:00 - 3:45)**
**[Screen: Application running at localhost:3000]**

**Actions:**
1. Show homepage with hero section
2. Click "Get Started" or navigate to Register
3. Fill registration form:
   - Name: "Demo Student"
   - Email: "demo@student.com"
   - Password: "password123"
   - Role: "Student"
   - Add dietary preferences: Vegetarian
   - Add allergies: Peanuts
4. Submit registration
5. Show success message

**Narration:**
"Users start at our responsive landing page. They can register with their details, select dietary preferences like vegetarian or vegan, and specify any allergies. This information helps filter menu items later."

---

#### **3.2 Login & Dashboard (3:45 - 4:15)**
**Actions:**
1. Navigate to Login page
2. Enter credentials: "demo@student.com" / "password123"
3. Show successful login
4. Display user menu in navbar
5. Show wallet balance display

**Narration:**
"After registration, users can log in. The system uses JWT tokens for authentication, validated at the API Gateway. Upon login, users see their profile in the navigation and their digital wallet balance."

---

#### **3.3 Menu Browsing (4:15 - 4:45)**
**Actions:**
1. Navigate to Menu page
2. Show menu items grid
3. Demonstrate category filters (Main, Beverages, Snacks, etc.)
4. Use search functionality
5. Show dietary tags and allergens on items
6. Add items to cart
7. Show cart preview

**Narration:**
"The menu page displays all available items with category filters and search functionality. Each item shows dietary tags like vegetarian or gluten-free, and allergen warnings. Users can add items to their cart, which persists in local storage."

---

#### **3.4 Cart & Checkout (4:45 - 5:30)**
**Actions:**
1. Navigate to Cart page
2. Show cart items with quantities
3. Adjust quantities
4. Select pickup date (tomorrow)
5. Select pickup time (12:30 PM)
6. Add special instructions
7. Show order summary with total
8. Click "Place Order"
9. Show order confirmation

**Narration:**
"In the cart, users review their selections, choose a convenient pickup date and time, and add any special instructions. The order total is calculated automatically. Upon placing the order, the system creates the order in the Order Service and deducts the amount from the digital wallet."

---

#### **3.5 Order Tracking (5:30 - 6:00)**
**Actions:**
1. Navigate to Orders page
2. Show list of orders
3. Show order status (pending)
4. Click on order to see details
5. Show order items, pickup time, total

**Narration:**
"Users can track their orders in real-time. Each order goes through a status workflow: pending, confirmed, preparing, ready, and completed. Users can also cancel pending orders if needed."

**[Transition to Admin Features]**

---

### **SECTION 4: Admin Features (6:00 - 7:30)**

#### **4.1 Admin Login (6:00 - 6:15)**
**Actions:**
1. Logout from student account
2. Login as admin: "admin@cafeteria.com" / "admin123"
3. Show admin dashboard

**Narration:**
"Now let's see the admin capabilities. Administrators have additional privileges to manage the system."

---

#### **4.2 Admin Dashboard (6:15 - 6:45)**
**Actions:**
1. Show dashboard statistics cards
2. Point out: Total Orders, Pending Orders, Revenue
3. Show quick action buttons

**Narration:**
"The admin dashboard provides an overview of system activity, including total orders, pending orders requiring attention, and revenue metrics. This helps canteen staff plan preparation and manage operations."

---

#### **4.3 Menu Management (6:45 - 7:15)**
**Actions:**
1. Navigate to Menu Management
2. Show existing menu items
3. Click "Add New Item"
4. Fill form:
   - Name: "Grilled Chicken Wrap"
   - Price: $10.99
   - Category: "Main"
   - Dietary tags: High Protein
   - Preparation time: 12 minutes
5. Save item
6. Show item added to menu

**Narration:**
"Admins can manage the entire menu through CRUD operations. They can add new items with details like price, category, dietary tags, and preparation time. They can also update availability and set daily order limits to manage inventory."

---

#### **4.4 Order Management (7:15 - 7:30)**
**Actions:**
1. Navigate to Orders
2. Show all orders
3. Find the order placed earlier
4. Update status from "pending" to "confirmed"
5. Show status update success

**Narration:**
"Admins can view all orders and update their status through the workflow. This keeps customers informed about their order progress."

**[Transition to Technical Summary]**

---

### **SECTION 5: Technical Implementation (7:30 - 8:30)**
**[Screen: Code/IDE or Architecture Diagram]**

**Narration:**
"Let me briefly highlight the technical implementation. The backend follows a microservices architecture with .NET 10, using Entity Framework Core with PostgreSQL. Each service has its own database following the database-per-service pattern."

**Show code snippets:**
1. **Microservices Structure:** Show folder structure
2. **API Gateway:** Show ocelot.json configuration
3. **JWT Authentication:** Show token validation
4. **Database Schema:** Show entity models
5. **Frontend:** Show React component structure

**Key Points:**
- Independent service deployment
- JWT authentication at Gateway
- Database per service
- React with TypeScript frontend
- Swagger API documentation

---

### **SECTION 6: Conclusion (8:30 - 9:00)**
**[Screen: Summary Slide]**

**Narration:**
"To summarize, the Cafeteria Pre-ordering System demonstrates:
- A complete full-stack application with microservices architecture
- JWT-based authentication and role-based access control
- Real-time order tracking and management
- Digital wallet integration
- Responsive UI for both users and admins

The system successfully addresses the problem of long cafeteria queues while providing a scalable, maintainable architecture."

---

### **OUTRO (9:00 - 9:30)**
**[Screen: Thank You / Contact Info]**

**Narration:**
"Thank you for watching this demonstration. The complete source code is available in the GitHub repository linked in the submission. Feel free to reach out with any questions."

---

## Technical Notes for Recording

### **Before Recording:**
1. Start all services in order:
   ```bash
   # Terminal 1: Auth Service
   cd microservices/AuthService && dotnet run --urls "http://localhost:5001"

   # Terminal 2: Menu Service
   cd microservices/MenuService && dotnet run --urls "http://localhost:5002"

   # Terminal 3: Order Service
   cd microservices/OrderService && dotnet run --urls "http://localhost:5003"

   # Terminal 4: API Gateway
   cd microservices/ApiGateway && dotnet run --urls "http://localhost:5010"

   # Terminal 5: Frontend
   cd frontend/cafeteria-client && npm start
   ```

2. Seed sample data:
   ```bash
   curl -X POST http://localhost:5010/api/menu/seed
   ```

3. Clear browser cache/history for clean demo

4. Prepare demo accounts:
   - Student: demo@student.com / password123
   - Admin: admin@cafeteria.com / admin123

### **Screen Recording Settings:**
- Resolution: 1920x1080 or 1440x900
- Frame rate: 30fps
- Audio: Clear narration
- Cursor: Visible and highlighted

### **Post-Production:**
1. Add title cards between sections
2. Zoom in on important UI elements
3. Add captions for key features
4. Background music (optional, low volume)
5. Export in 1080p quality

---

## Demo Checklist

### **Pre-Demo:**
- [ ] All services running
- [ ] Database seeded with sample data
- [ ] Browser open at localhost:3000
- [ ] Demo accounts ready
- [ ] Screen recorder configured

### **During Demo:**
- [ ] Speak clearly and at moderate pace
- [ ] Show mouse movements
- [ ] Wait for page loads
- [ ] Highlight key features
- [ ] Show error handling (optional)

### **Post-Demo:**
- [ ] Review recording quality
- [ ] Check audio levels
- [ ] Verify all sections covered
- [ ] Upload to Google Drive
- [ ] Set sharing permissions (BITS emails)

---

**End of Script**
"""

    output_path = '/Users/subhodipanwesa/Documents/BITS_WILP/FullStack/Assignment1/Demo_Video_Script.md'
    with open(output_path, 'w') as f:
        f.write(script)
    print(f"✅ Demo video script created: {output_path}")
    return output_path


if __name__ == '__main__':
    print("=" * 60)
    print("Creating Documentation Files for FSAD Assignment")
    print("=" * 60)
    print()

    # Create combined documentation
    print("Step 1: Creating Combined Architecture Documentation...")
    create_combined_documentation()
    print()

    # Create API documentation
    print("Step 2: Creating API Documentation...")
    create_api_documentation()
    print()

    # Create demo script
    print("Step 3: Creating Demo Video Script...")
    create_demo_script()
    print()

    print("=" * 60)
    print("All documents created successfully!")
    print("=" * 60)
    print()
    print("Files created:")
    print("  1. Complete_Architecture_Documentation.docx")
    print("  2. API_Documentation.docx")
    print("  3. Demo_Video_Script.md")
    print()
    print("These files are located in:")
    print("  /Users/subhodipanwesa/Documents/BITS_WILP/FullStack/Assignment1/")
